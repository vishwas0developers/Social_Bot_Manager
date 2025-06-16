# app.py (Complete Code with Chunking, Custom Prompts, Cleaning, Revised Excel)

import os
import uuid
import shutil
import zipfile
import json
import pandas as pd
import unicodedata # <<< Required for cleaning function
from docx import Document
from flask import (
    Flask, request, render_template, send_file,
    redirect, url_for, flash, current_app, jsonify
)
from werkzeug.utils import secure_filename
import traceback
import re # For option parsing and sorting key
import time # For potential delays in retries
from markupsafe import escape # Import escape for sanitizing user input
from utils.cleaners import clean_excel_string, clean_docx_string
# --- Utils Import ---
# Ensure this points to the correct function implementing the chunking strategy
try:
    from utils.ocr_processor import (
        pdf_to_images,
        call_ocr_model_in_chunks, # <<< Use the chunking function
        SUPPORTED_MODELS,
        get_gemini_api_key_state,
        default_gemini_model_name, # For setting default in UI
        build_raw_ocr_prompt,
        call_gemini_raw_text,
        _load_images_for_prompt
    )
except ImportError as e:
    print(f"❌ Error importing from utils: {e}. Make sure utils/ocr_processor.py and utils/__init__.py exist and define 'call_ocr_model_in_chunks'.")
    exit()

# --- Basic Flask App Setup ---
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", os.urandom(24))

# --- Folder Configuration ---
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 30 * 1024 * 1024 # Keep 30MB limit
os.makedirs(UPLOAD_FOLDER, exist_ok=True); os.makedirs(OUTPUT_FOLDER, exist_ok=True)
print(f"📂 Upload folder: {app.config['UPLOAD_FOLDER']}")
print(f"📂 Output folder: {app.config['OUTPUT_FOLDER']}")

# --- Helper Functions ---
def allowed_file(filename):
    """Checks if the file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_question_sort_key(question_object):
    """Extracts a sortable key (integer) from the question number string."""
    q_num_str = question_object.get("question_number", "")
    if not q_num_str or not isinstance(q_num_str, str): return float('inf')
    match = re.search(r'\d+', q_num_str)
    if match:
        try: return int(match.group(0))
        except ValueError: return float('inf')
    else: return float('inf')

# --- Routes ---

@app.route('/')
def index():
    """Homepage (upload form) - Pass supported models, API key status, and default model."""
    gemini_api_key_ok = get_gemini_api_key_state()
    models_to_show = SUPPORTED_MODELS

    default_model_to_pass = None
    gemini_models_list = models_to_show.get('gemini', [])
    if gemini_api_key_ok and default_gemini_model_name in gemini_models_list:
         default_model_to_pass = default_gemini_model_name

    return render_template('index.html',
                           supported_models=models_to_show,
                           gemini_api_key_ok=gemini_api_key_ok,
                           default_model=default_model_to_pass
                           )

@app.route('/upload', methods=['POST'])
def upload_file():
    start_time = time.time()
    job_upload_path, job_temp_image_path, temp_images_created = None, None, False # Initialize for finally block
    # Initialize job_env earlier for potential use in except block if needed
    job_env = {"job_id": "UNKNOWN"} # Default if setup fails early

    try:
        # 1. Parse and Validate Request
        params, error_response = _parse_and_validate_request(request)
        if error_response:
            return error_response # Returns a redirect with flash message

        # --- Add Safety Check (Recommended from previous step, keeping it) ---
        # This guards against the observed KeyError if validation logic has subtle issues.
        if not params or 'original_filename' not in params:
            print("❌ Internal Error: Params dictionary missing or 'original_filename' key absent after validation returned success.")
            flash("An internal server error occurred (invalid request state). Please try uploading again.", "error")
            return redirect(url_for('index'))
        # --- End Safety Check ---

        # 2. Setup Job Environment
        # Now we are sure params exists and has the key
        job_env = _setup_job_environment(params['original_filename'])
        job_upload_path = job_env['job_upload_path'] # Needed for cleanup
        job_temp_image_path = job_env['job_temp_image_path'] # Needed for cleanup

        # 3. Prepare Image Inputs (Saves upload, converts PDF)
        image_paths, temp_images_created, error_response = _prepare_image_inputs(
            params['is_pdf'],
            job_env['job_upload_path'],
            job_env['job_temp_image_path']
        )
        if error_response:
             # _prepare_image_inputs should ideally handle its own cleanup on failure
             # Cleanup might still happen in the 'finally' block below too
             return error_response

        # 4. Process based on Mode
        num_questions = 0 # Default
        processing_successful = False

        # --- UPDATED PROCESSING LOGIC (from target code) ---
        if params['content_mode'] == 'raw':
            print("Routing to: _process_raw_mode")
            processing_successful = _process_raw_mode(
                job_output_path=job_env['job_output_path'],
                output_basename=job_env['output_basename'],
                image_paths_to_process=image_paths,
                language=params['language'],
                model=params['model'],
                custom_prompt=params['custom_prompt']
            )
            # num_questions remains 0 for raw mode

        elif params['content_mode'] == 'raw_math': # <<< NEW BRANCH from target
            print("Routing to: _process_raw_math_mode")
            processing_successful = _process_raw_math_mode( # <<< CALL NEW FUNCTION from target
                job_output_path=job_env['job_output_path'],
                output_basename=job_env['output_basename'],
                image_paths_to_process=image_paths,
                language=params['language'],
                model=params['model'],
                custom_prompt=params['custom_prompt']
            )
            # num_questions remains 0 for raw math mode

        else: # MCQ Mode (Text or Math)
            print(f"Routing to: _process_mcq_mode (Mode: {params['content_mode']})")
            processing_successful, num_questions = _process_mcq_mode(
                job_output_path=job_env['job_output_path'],
                output_basename=job_env['output_basename'],
                image_paths_to_process=image_paths,
                company=params['company'],
                model=params['model'],
                language=params['language'],
                content_mode=params['content_mode'], # Pass 'text' or 'math' here
                generate_explanation=params['generate_explanation'],
                generate_answer=params['generate_answer'],
                custom_prompt=params['custom_prompt'],
                is_pdf=params['is_pdf'] # Needed for ZIP generation inside _process_mcq_mode
            )
        # ----------------------------------------------------

        # 5. Redirect to Results if Successful
        if processing_successful:
            # Pass necessary info derived from 'params' and processing results
            # (This part matches the target code structure)
            return redirect(url_for('show_results',
                job_id=job_env['job_id'],
                is_pdf=params['is_pdf'],
                language=params['language'],
                company=params['company'],
                model=params['model'],
                content_mode=params['content_mode'], # <<< Pass 'raw', 'raw_math', 'text', or 'math'
                generate_explanation=params['generate_explanation'],
                generate_answer=params['generate_answer'],
                custom_prompt_used=bool(params['custom_prompt']),
                num_pages=len(image_paths),
                num_questions=num_questions
            ))
        else:
             # Processing functions should flash specific errors if possible
             flash("Processing failed. Please check logs or try again.", "error")
             # Cleanup happens in finally, redirect to index
             return redirect(url_for('index'))

    except Exception as e:
        # Use the job_id from job_env if it was set, otherwise default
        job_id_for_log = job_env.get('job_id', 'UNKNOWN')
        print(f"❌ An unexpected error occurred in upload_file orchestrator for job {job_id_for_log}:")
        traceback.print_exc() # Print the full traceback for debugging
        flash(f"Unexpected server error occurred: {e}", "error")
        # Fall through to finally block for cleanup, then redirect
        return redirect(url_for('index'))

    finally:
        # 6. Cleanup
        # Ensure paths exist before trying to clean up
        _cleanup_job_files(job_upload_path, job_temp_image_path, temp_images_created)
        end_time = time.time()
        # Use potentially updated job_id from job_env
        job_id_for_log = job_env.get('job_id', 'UNKNOWN')
        print(f"⏱️ Total request handling for job {job_id_for_log} took {end_time - start_time:.2f} seconds.")

# --- Define Helper Functions ---

def _parse_and_validate_request(request):
    """Parses form data and file, validates inputs. Ensures original_filename is set early."""
    params = {}
    error_response = None

    # 1. Basic File Presence Check
    if 'file' not in request.files:
        flash('No file selected', 'error')
        return None, redirect(url_for('index')) # Return tuple (None, redirect)
    file = request.files['file']

    # 2. Check for Empty Filename from Browser
    if file.filename == '':
        flash('Empty filename', 'error')
        return None, redirect(url_for('index')) # Return tuple (None, redirect)

    # 3. --- Secure and Assign Filename EARLY ---
    # Secure the filename provided by the browser.
    s_filename = secure_filename(file.filename)
    # Check if secure_filename resulted in empty (e.g., filename was just "..")
    if not s_filename:
         flash('Invalid filename provided.', 'error')
         return None, redirect(url_for('index')) # Return tuple (None, redirect)

    # Assign the secured filename to params. Now the key definitely exists if we proceed.
    params['original_filename'] = s_filename
    params['file'] = file # Store file object if needed (though often re-accessed via request)
    # ------------------------------------------

    # 4. File Type Check (using the original name from 'file' object for extension)
    if not allowed_file(file.filename):
        flash(f'Invalid file type. Allowed: {", ".join(ALLOWED_EXTENSIONS)}', 'error')
        return None, redirect(url_for('index')) # Return tuple (None, redirect)

    # 5. Set Derived Parameters (Now guaranteed original_filename exists)
    params['is_pdf'] = params['original_filename'].lower().endswith('.pdf')

    # 6. Parse Form Data
    params['language'] = request.form.get('language', 'English')
    params['company'] = request.form.get('company', 'gemini')
    params['model'] = request.form.get('model')
    params['content_mode'] = request.form.get('content_mode', 'text')
    params['generate_explanation'] = request.form.get('generate_explanation') == 'on'
    params['generate_answer'] = request.form.get('generate_answer') == 'on'
    params['raw_mode'] = params['content_mode'] in ['raw', 'raw_math']
    custom_prompt_raw = request.form.get('custom_prompt', '')
    params['custom_prompt'] = str(escape(custom_prompt_raw)).strip() if custom_prompt_raw else None

    # 7. Form Validation Checks
    if not params['company'] or params['company'] not in SUPPORTED_MODELS:
        flash('Invalid Company selected.', 'error')
        return None, redirect(url_for('index')) # Return tuple (None, redirect)
    if not params['model'] or params['model'] not in SUPPORTED_MODELS.get(params['company'], []):
        flash(f'Invalid Model selected for {params["company"]}.', 'error')
        return None, redirect(url_for('index')) # Return tuple (None, redirect)
    if params['company'] == 'gemini' and not get_gemini_api_key_state():
        flash('Gemini API Key missing or invalid. Cannot process with Gemini.', 'error')
        return None, redirect(url_for('index')) # Return tuple (None, redirect)

    # 8. Success: All checks passed, return params and no error response
    print(f"✅ Validation successful. Params keys: {list(params.keys())}") # Add log for confirmation
    return params, None

def _setup_job_environment(original_filename):
    """Sets up job ID and paths."""
    job_id = str(uuid.uuid4())
    job_output_path = os.path.join(current_app.config['OUTPUT_FOLDER'], job_id)
    job_upload_path = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{job_id}_{original_filename}")
    job_temp_image_path = os.path.join(job_output_path, 'temp_images')
    output_basename = "ocr_output"
    os.makedirs(job_output_path, exist_ok=True)
    # Also create temp path dir structure immediately if needed? Or let pdf_to_images handle it.
    # os.makedirs(job_temp_image_path, exist_ok=True) # Optional

    return {
        "job_id": job_id,
        "job_output_path": job_output_path,
        "job_upload_path": job_upload_path,
        "job_temp_image_path": job_temp_image_path,
        "output_basename": output_basename
    }

def _prepare_image_inputs(is_pdf, job_upload_path, job_temp_image_path):
    """Saves the uploaded file and converts PDF to images if needed."""
    image_paths_to_process = []
    temp_images_created = False
    error_response = None

    try:
        # Access the file object from the request context (safer than passing it)
        file = request.files['file']
        file.save(job_upload_path)
        print(f"📂 File saved to: {job_upload_path}")

        if is_pdf:
            print(f"🔄 Converting PDF to images...")
            image_paths_to_process = pdf_to_images(job_upload_path, job_temp_image_path)
            if image_paths_to_process:
                 temp_images_created = True
                 print(f"   ✅ Generated {len(image_paths_to_process)} images.")
            else:
                 flash("Failed to convert PDF to images.", "error")
                 # Try cleanup here before returning error response
                 if os.path.exists(job_temp_image_path): shutil.rmtree(job_temp_image_path)
                 if os.path.exists(job_upload_path): os.remove(job_upload_path)
                 error_response = redirect(url_for('index'))
        else:
            # Use the saved image path directly
            image_paths_to_process.append(job_upload_path)
            print(f"   Processing single image: {job_upload_path}")

    except Exception as e:
        print(f"❌ Error during file save or PDF conversion: {e}")
        traceback.print_exc()
        flash(f"Error handling uploaded file: {e}", "error")
        error_response = redirect(url_for('index'))
        # Ensure potential partial files are cleaned up if possible? (Difficult without knowing state)

    return image_paths_to_process, temp_images_created, error_response

def _process_raw_mode(job_output_path, output_basename, image_paths_to_process, language, model, custom_prompt):
    """Handles the processing logic for Raw Text mode."""
    print(f"⚙️ Executing Raw Mode Logic...")
    doc_output_path = os.path.join(job_output_path, f"{output_basename}.docx")
    raw_text_output_path = os.path.join(job_output_path, f"{output_basename}_raw.txt")

    try:
        doc = Document()
        doc.add_heading('Extracted OCR Text (Raw Mode - Structure Attempt)', level=1)
        all_raw_text_for_preview = ""

        for i, image_path in enumerate(image_paths_to_process):
            page_num_for_log = i + 1
            print(f"   📄 Processing Page {page_num_for_log} (Raw)...")
            loaded_images = _load_images_for_prompt([image_path])
            doc.add_heading(f'Page {page_num_for_log}', level=2) # Add page heading regardless

            if not loaded_images:
                print(f"      ⚠️ Skipping page {page_num_for_log} due to image loading failure.")
                run = doc.add_paragraph().add_run("(Failed to load image for this page)")
                run.italic = True
                all_raw_text_for_preview += f"\n\n--- Page {page_num_for_log} ---\n\n(Failed to load image)\n"
                continue

            prompt = build_raw_ocr_prompt(language, custom_prompt)
            parsed_raw_page_text = call_gemini_raw_text(model, [prompt], loaded_images) # Simplified call

            if parsed_raw_page_text and isinstance(parsed_raw_page_text, str):
                raw_text_page = parsed_raw_page_text.strip()
                all_raw_text_for_preview += f"\n\n--- Page {page_num_for_log} ---\n\n{raw_text_page}\n"
                print(f"      ✅ Extracted text for page {page_num_for_log} ({len(raw_text_page)} chars).")

                text_blocks = raw_text_page.replace('\r\n', '\n').split('\n\n')
                added_paragraph_to_docx = False
                for block in text_blocks:
                    cleaned_block = block.strip()
                    if cleaned_block:
                        doc.add_paragraph(cleaned_block)
                        added_paragraph_to_docx = True
                if not added_paragraph_to_docx:
                     run = doc.add_paragraph().add_run("(No text content blocks found after processing for this page)")
                     run.italic = True
                     print(f"      ℹ️ No non-empty text blocks added to DOCX for page {page_num_for_log}.")
            else:
                print(f"      ⚠️ Failed to extract text from page {page_num_for_log}.")
                run = doc.add_paragraph().add_run("(Failed to extract text from this page)")
                run.italic = True
                all_raw_text_for_preview += f"\n\n--- Page {page_num_for_log} ---\n\n(Failed to extract text)\n"

        # Save TXT Preview
        if all_raw_text_for_preview:
            try:
                with open(raw_text_output_path, 'w', encoding='utf-8') as f: f.write(all_raw_text_for_preview.strip())
                print(f"✅ Raw text preview file saved to: {raw_text_output_path}")
            except Exception as write_err:
                print(f"❌ Error writing raw text preview file: {write_err}")
                flash("Warning: Could not save the text preview file.", "warning") # Non-fatal
        else:
            print("⚠️ No text was accumulated, skipping raw preview file generation.")

        # Save DOCX
        doc.save(doc_output_path)
        print(f"✅ Raw DOCX (structure attempted) generated and saved to: {doc_output_path}")
        return True # Indicate success

    except Exception as e:
        print(f"❌ Error during Raw Mode processing: {e}")
        traceback.print_exc()
        flash(f"Error during raw text processing: {e}", "error")
        # Attempt to save DOCX even if partial? Maybe not.
        return False # Indicate failure

def _process_raw_math_mode(job_output_path, output_basename, image_paths_to_process, language, model, custom_prompt):
    """Handles the processing logic for Raw Text mode WITH LaTeX math formatting."""
    print(f"⚙️ Executing Raw Math Mode Logic (Requesting LaTeX)...") # <<< Note the change
    doc_output_path = os.path.join(job_output_path, f"{output_basename}.docx")
    # Using a distinct name for the raw text preview in this mode might be clearer
    raw_text_output_path = os.path.join(job_output_path, f"{output_basename}_raw_math.txt")

    try:
        doc = Document()
        doc.add_heading('Extracted OCR Text (Raw Mode - Math/LaTeX Attempt)', level=1) # <<< Heading change
        all_raw_text_for_preview = ""

        for i, image_path in enumerate(image_paths_to_process):
            page_num_for_log = i + 1
            print(f"   📄 Processing Page {page_num_for_log} (Raw Math)...") # <<< Log change
            loaded_images = _load_images_for_prompt([image_path])
            doc.add_heading(f'Page {page_num_for_log}', level=2)

            if not loaded_images:
                print(f"      ⚠️ Skipping page {page_num_for_log} due to image loading failure.")
                run = doc.add_paragraph().add_run("(Failed to load image for this page)")
                run.italic = True
                all_raw_text_for_preview += f"\n\n--- Page {page_num_for_log} ---\n\n(Failed to load image)\n"
                continue

            # --- CRITICAL CHANGE: Use the new prompt builder ---
            from utils.ocr_processor import build_raw_math_ocr_prompt # Ensure imported
            prompt = build_raw_math_ocr_prompt(language, custom_prompt)
            # --- Use the same raw text API call function ---
            parsed_raw_page_text = call_gemini_raw_text(model, [prompt], loaded_images) # Reuse existing raw call

            if parsed_raw_page_text and isinstance(parsed_raw_page_text, str):
                raw_text_page = parsed_raw_page_text.strip()
                all_raw_text_for_preview += f"\n\n--- Page {page_num_for_log} ---\n\n{raw_text_page}\n"
                print(f"      ✅ Extracted text for page {page_num_for_log} ({len(raw_text_page)} chars, LaTeX requested).") # <<< Log change

                # --- Logic for adding to DOCX remains the same ---
                # It will just contain the text as returned, including \(...\)
                text_blocks = raw_text_page.replace('\r\n', '\n').split('\n\n')
                added_paragraph_to_docx = False
                for block in text_blocks:
                    cleaned_block = block.strip()
                    if cleaned_block:
                        doc.add_paragraph(cleaned_block)
                        added_paragraph_to_docx = True
                if not added_paragraph_to_docx:
                     run = doc.add_paragraph().add_run("(No text content blocks found after processing for this page)")
                     run.italic = True
                     print(f"      ℹ️ No non-empty text blocks added to DOCX for page {page_num_for_log}.")
            else:
                print(f"      ⚠️ Failed to extract text from page {page_num_for_log} (Raw Math Mode).") # <<< Log change
                run = doc.add_paragraph().add_run("(Failed to extract text from this page)")
                run.italic = True
                all_raw_text_for_preview += f"\n\n--- Page {page_num_for_log} ---\n\n(Failed to extract text)\n"

        # Save TXT Preview (with potentially raw LaTeX)
        if all_raw_text_for_preview:
            try:
                # Save to the specific file name
                with open(raw_text_output_path, 'w', encoding='utf-8') as f: f.write(all_raw_text_for_preview.strip())
                print(f"✅ Raw math text preview file saved to: {raw_text_output_path}") # <<< Log change
            except Exception as write_err:
                print(f"❌ Error writing raw math text preview file: {write_err}")
                flash("Warning: Could not save the text preview file.", "warning")
        else:
            print("⚠️ No text was accumulated, skipping raw math preview file generation.") # <<< Log change

        # Save DOCX (will contain the raw text with \(...\))
        doc.save(doc_output_path)
        print(f"✅ Raw Math DOCX (LaTeX attempted) generated and saved to: {doc_output_path}") # <<< Log change
        return True

    except Exception as e:
        print(f"❌ Error during Raw Math Mode processing: {e}") # <<< Log change
        traceback.print_exc()
        flash(f"Error during raw math text processing: {e}", "error")
        return False
    
def _process_mcq_mode(job_output_path, output_basename, image_paths_to_process, company, model, language, content_mode, generate_explanation, generate_answer, custom_prompt, is_pdf):
    """Handles the processing logic for MCQ Text/Math mode."""
    print(f"⚙️ Executing MCQ Mode Logic ({content_mode})...")
    final_sorted_questions = []
    num_questions = 0
    json_output_path = os.path.join(job_output_path, f"{output_basename}.json")
    excel_output_path = os.path.join(job_output_path, f"{output_basename}.xlsx")
    doc_output_path = os.path.join(job_output_path, f"{output_basename}.docx")
    zip_output_path = os.path.join(job_output_path, f"{output_basename}_structured.zip")

    try:
        # --- Call the Chunking OCR Processor ---
        parsed_output_data = call_ocr_model_in_chunks(
            company=company, model=model, all_image_paths=image_paths_to_process,
            language=language, content_mode=content_mode,
            generate_explanation=generate_explanation, generate_answer=generate_answer,
            custom_prompt=custom_prompt, max_retries_per_chunk=2,
            chunk_size=1, retry_delay_seconds=5
        )

        # --- Validate OCR Output ---
        if not parsed_output_data or not isinstance(parsed_output_data, dict):
            error_message = "MCQ processing failed or returned invalid data."
            if isinstance(parsed_output_data, dict) and "_processing_warnings" in parsed_output_data:
                warnings = parsed_output_data["_processing_warnings"]
                error_message += f" Details: {', '.join(warnings)}"
            flash(error_message, "error")
            return False, 0 # Failure

        # --- Grouping & Sorting MCQs ---
        page_keys = sorted(
            [k for k in parsed_output_data if k.startswith("page_") and k.endswith("_MCQs")],
            key=lambda x: int(x.split('_')[1])
        )
        for page_key in page_keys:
            try:
                page_num = int(page_key.split('_')[1])
                questions_on_page = parsed_output_data.get(page_key, [])
                if isinstance(questions_on_page, list):
                    for q in questions_on_page:
                        if isinstance(q, dict):
                            q['original_page'] = q.get('original_page', page_num)
                            final_sorted_questions.append(q)
                else: print(f"⚠️ Warning: Data for key '{page_key}' was not a list, skipping.")
            except (ValueError, IndexError): print(f"⚠️ Warning: Could not parse page number from key '{page_key}', skipping.")

        final_sorted_questions.sort(key=get_question_sort_key)
        num_questions = len(final_sorted_questions)
        print(f"📊 Extracted and sorted {num_questions} questions.")

        # --- Generate JSON Output ---
        # (Error handling for individual file saves is kept internal to this function)
        final_json_output = {
            key: value for key, value in parsed_output_data.items()
            if (key.startswith("page_") and key.endswith("_MCQs")) or key == "header_footer_info"
        }
        if "_processing_warnings" in parsed_output_data:
            final_json_output["processing_info"] = {"chunk_warnings": parsed_output_data["_processing_warnings"]}
        try:
            with open(json_output_path, 'w', encoding='utf-8') as f: json.dump(final_json_output, f, ensure_ascii=False, indent=4)
            print(f"✅ JSON output saved to: {json_output_path}")
        except Exception as json_save_err: print(f"❌ Error saving JSON file: {json_save_err}") # Log and continue

        # --- Generate Excel Output ---
        if final_sorted_questions:
            try:
                excel_rows = []
                for idx, q in enumerate(final_sorted_questions):
                    q_cleaned = {k: clean_excel_string(v) if isinstance(v, str) else v for k, v in q.items()}
                    options_cleaned = [clean_excel_string(opt) for opt in q_cleaned.get("options", [])]
                    excel_rows.append({
                        "Sr. No": idx + 1, "Question Number": q_cleaned.get("question_number", ""),
                        "Original Page": q_cleaned.get("original_page", ""), "Question Text": q_cleaned.get("question_text", ""),
                        "Option A": options_cleaned[0] if len(options_cleaned) > 0 else "", "Option B": options_cleaned[1] if len(options_cleaned) > 1 else "",
                        "Option C": options_cleaned[2] if len(options_cleaned) > 2 else "", "Option D": options_cleaned[3] if len(options_cleaned) > 3 else "",
                        "Option E": options_cleaned[4] if len(options_cleaned) > 4 else "", "Answer": q_cleaned.get("answer", None),
                        "Explanation": q_cleaned.get("explanation", None), "Source Info": q_cleaned.get("source_info", "")
                    })
                df = pd.DataFrame(excel_rows, columns=["Sr. No", "Question Number", "Original Page", "Question Text", "Option A", "Option B", "Option C", "Option D", "Option E", "Answer", "Explanation", "Source Info"])
                df.to_excel(excel_output_path, index=False, engine='openpyxl')
                print(f"✅ Excel output saved to: {excel_output_path}")
            except Exception as excel_save_err: print(f"❌ Error saving Excel file: {excel_save_err}") # Log and continue
        else: print("ℹ️ No questions extracted, skipping Excel generation.")

        # --- Generate DOCX Output (for MCQs) ---
        if final_sorted_questions:
            try:
                doc_mcq = Document()
                doc_mcq.add_heading('Extracted Questions', level=1)
                for q in final_sorted_questions:
                    # (Cleaning logic remains the same)
                    q_num_clean=clean_docx_string(q.get('question_number','?')); q_text_clean=clean_docx_string(q.get('question_text','')); opts_clean=[clean_docx_string(opt) for opt in q.get('options',[])]; ans_clean=clean_docx_string(q.get('answer')) if q.get('answer') is not None else None; exp_clean=clean_docx_string(q.get('explanation')) if q.get('explanation') is not None else None; src_clean=clean_docx_string(q.get('source_info',''))
                    doc_mcq.add_heading(f"Q. {q_num_clean}", level=2)
                    if q_text_clean: doc_mcq.add_paragraph(q_text_clean)
                    for i, opt in enumerate(opts_clean): doc_mcq.add_paragraph(f"{chr(ord('A') + i)}) {opt}")
                    if ans_clean is not None: doc_mcq.add_paragraph(f"Answer: {ans_clean}")
                    if exp_clean is not None: doc_mcq.add_paragraph(f"Explanation: {exp_clean}")
                    if src_clean: doc_mcq.add_paragraph(f"(Source: {src_clean})")
                    doc_mcq.add_paragraph() # Spacing
                doc_mcq.save(doc_output_path)
                print(f"✅ DOCX output (MCQ) saved to: {doc_output_path}")
            except Exception as doc_save_err: print(f"❌ Error saving MCQ DOCX file: {doc_save_err}") # Log and continue
        else: print("ℹ️ No questions extracted, skipping DOCX generation (MCQ mode).")

        # --- Generate ZIP Output ---
        if is_pdf:
            files_to_zip = []
            if os.path.exists(json_output_path): files_to_zip.append((json_output_path, f"{output_basename}.json"))
            if os.path.exists(excel_output_path): files_to_zip.append((excel_output_path, f"{output_basename}.xlsx"))
            if os.path.exists(doc_output_path): files_to_zip.append((doc_output_path, f"{output_basename}.docx"))
            if files_to_zip:
                try:
                    with zipfile.ZipFile(zip_output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for path, arcname in files_to_zip: zipf.write(path, arcname=arcname)
                    print(f"✅ ZIP archive created: {zip_output_path}")
                except Exception as zip_err: print(f"❌ Error creating ZIP file: {zip_err}") # Log and continue
            else: print("ℹ️ No structured output files generated to zip.")

        return True, num_questions # Success

    except Exception as e:
        print(f"❌ Error during MCQ Mode processing: {e}")
        traceback.print_exc()
        flash(f"Error during MCQ processing: {e}", "error")
        return False, 0 # Failure

def _cleanup_job_files(job_upload_path, job_temp_image_path, temp_images_created):
    """Cleans up temporary files generated during processing."""
    print("🧹 Starting cleanup...")
    # Remove uploaded file (unless it was a direct image upload handled differently - check logic in orchestrator if needed)
    if job_upload_path and os.path.exists(job_upload_path):
         # Add logic here if direct image uploads should NOT be deleted
         try:
             # Check if it was a PDF conversion source or just an upload to discard
             # Assuming we always discard the file copied to job_upload_path for now
             os.remove(job_upload_path)
             print(f"🧹 Cleaned up uploaded file: {job_upload_path}")
         except OSError as remove_err:
             print(f"⚠️ Warning: Could not remove uploaded file {job_upload_path}. Error: {remove_err}")

    # Remove temp image folder ONLY if it was created (i.e., from a PDF)
    if temp_images_created and job_temp_image_path and os.path.exists(job_temp_image_path):
        try:
            shutil.rmtree(job_temp_image_path)
            print(f"🧹 Cleaned up temp image directory: {job_temp_image_path}")
        except OSError as rmtree_err:
            print(f"⚠️ Warning: Could not remove temp image directory {job_temp_image_path}. Error: {rmtree_err}")
    print("🧹 Cleanup finished.")

def deep_clean_dict(obj):
    if isinstance(obj, dict):
        return {k: deep_clean_dict(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [deep_clean_dict(item) for item in obj]
    elif isinstance(obj, str):
        return clean_excel_string(obj)
    else:
        return obj

@app.route('/results/<job_id>')
def show_results(job_id):
    """
    Displays the results page, loading either raw text preview or structured MCQ data.
    Includes fix for loading raw and raw_math text preview files.
    """
    # --- Get Parameters from Query String ---
    is_pdf = request.args.get('is_pdf', 'False').lower() == 'true'
    language = request.args.get('language', 'N/A')
    company = request.args.get('company', 'N/A')
    model = request.args.get('model', 'N/A')
    content_mode = request.args.get('content_mode', 'text')
    generate_explanation = request.args.get('generate_explanation', 'false').lower() == 'true'
    generate_answer = request.args.get('generate_answer', 'false').lower() == 'true'
    custom_prompt_used = request.args.get('custom_prompt_used', 'false').lower() == 'true'
    num_pages = request.args.get('num_pages', 0, type=int)
    num_questions = request.args.get('num_questions', 0, type=int)

    # --- Construct File Paths ---
    safe_job_id = str(job_id).replace("..", "").replace("/", "")
    if safe_job_id != job_id:
        flash("Invalid job ID detected.", "error")
        return redirect(url_for('index'))

    job_output_path = os.path.join(current_app.config['OUTPUT_FOLDER'], safe_job_id)
    output_basename = "ocr_output"
    json_filename = f"{output_basename}.json"
    json_path = os.path.join(job_output_path, json_filename)
    raw_text_filename = f"{output_basename}_raw.txt"
    raw_text_path = os.path.join(job_output_path, raw_text_filename)
    raw_math_text_filename = f"{output_basename}_raw_math.txt"
    raw_math_text_path = os.path.join(job_output_path, raw_math_text_filename)

    print(f"📊 Showing results for job_id: {safe_job_id} (Mode: {content_mode}, Custom Prompt: {custom_prompt_used})")

    # --- Initialize ---
    grouped_data_for_template = {}
    chunk_warnings = []
    raw_text_preview = ""

    # --- Check if Output Directory Exists ---
    if not os.path.isdir(job_output_path):
        print(f"   ❌ Output directory not found: {job_output_path}")
        flash(f"Results not found for job {safe_job_id}.", "error")
        return redirect(url_for('index'))

    # --- Load Data Based on Mode ---
    if content_mode in ['raw', 'raw_math']:
        preview_file_path = raw_math_text_path if content_mode == 'raw_math' else raw_text_path
        mode_label = "Raw Math (LaTeX)" if content_mode == 'raw_math' else "Raw Text"
        print(f"   Attempting to load {mode_label} preview text from: {preview_file_path}")

        if os.path.exists(preview_file_path):
            try:
                with open(preview_file_path, 'r', encoding='utf-8') as f:
                    raw_text_preview = f.read()
                if raw_text_preview:
                    print(f"   ✅ Loaded {mode_label} preview text ({len(raw_text_preview)} chars).")
                else:
                    print(f"   ⚠️ {mode_label} preview is empty.")
                    raw_text_preview = f"[{mode_label} Preview file is empty]"
            except Exception as read_err:
                print(f"   ❌ Error reading {mode_label} preview: {read_err}")
                raw_text_preview = f"[Error reading {mode_label} preview: {read_err}]"
        else:
            print(f"   ⚠️ {mode_label} preview not found at: {preview_file_path}")
            raw_text_preview = f"[{mode_label} Preview file not found]"
            flash(f"{mode_label} preview not found. Download the DOCX for full content.", "warning")

    elif content_mode in ['text', 'math']:
        print(f"   Attempting to load structured MCQ data from: {json_path}")
        if not os.path.exists(json_path):
            print(f"   ❌ JSON file missing at: {json_path}")
            flash("Structured MCQ data not found. Try downloading the available files.", "error")
        else:
            try:
                with open(json_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

                grouped_data_for_template = {
                    k: v for k, v in data.items()
                    if isinstance(k, str) and k.startswith("page_") and k.endswith("_MCQs")
                }

                grouped_data_for_template = deep_clean_dict(grouped_data_for_template)
                chunk_warnings = data.get("processing_info", {}).get("chunk_warnings", [])

                if chunk_warnings:
                    print(f"   ⚠️ Warnings during processing: {chunk_warnings}")
                    flash("Some parts failed to process properly. See warnings below.", "warning")

                if grouped_data_for_template:
                    print(f"   ✅ Loaded MCQ data for {len(grouped_data_for_template)} pages.")
                else:
                    print("   ℹ️ JSON found but contains no MCQ data.")
                    flash("No MCQs were found in the output JSON.", "info")

            except json.JSONDecodeError as json_err:
                print(f"   ❌ JSON decoding error: {json_err}")
                flash("Output file is corrupted and can't be displayed.", "error")
            except Exception as e:
                print(f"   ❌ Unexpected error reading results JSON: {e}")
                traceback.print_exc()
                flash("Error occurred while loading structured result data.", "error")
        raw_text_preview = ""

    else:
        print(f"   ⚠️ Invalid content_mode: {content_mode}")
        flash(f"Unknown mode '{content_mode}' detected.", "warning")
        raw_text_preview = f"[Unknown processing mode: {content_mode}]"

    return render_template('results.html',
                           job_id=safe_job_id,
                           is_pdf=is_pdf,
                           language=language,
                           company=company,
                           model=model,
                           content_mode=content_mode,
                           generate_explanation=generate_explanation,
                           generate_answer=generate_answer,
                           custom_prompt_used=custom_prompt_used,
                           num_pages=num_pages,
                           num_questions=num_questions,
                           raw_text_preview=raw_text_preview,
                           grouped_questions_json=json.dumps(grouped_data_for_template),
                           chunk_warnings=chunk_warnings)

@app.route('/download/<job_id>/<file_format>')
def download_file(job_id, file_format):
    """Serves generated files for download."""
    job_output_path = os.path.join(current_app.config['OUTPUT_FOLDER'], job_id); output_basename = "ocr_output"
    filename_map = {
        'json': f"{output_basename}.json",
        'excel': f"{output_basename}.xlsx",
        'doc': f"{output_basename}.docx",
        'zip': f"{output_basename}_structured.zip"
    }
    if file_format not in filename_map: flash('Invalid file format.', 'error'); return redirect(url_for('show_results', job_id=job_id))

    filename = filename_map[file_format];
    file_path = os.path.join(job_output_path, filename)

    print(f"⬇️ Attempting download: {file_path}")
    if not os.path.exists(file_path):
        print(f"   ❌ File not found: {file_path}")
        if not os.path.exists(job_output_path):
             print(f"   ❌ Output dir missing: {job_output_path}")
             flash(f'Output folder missing for job {job_id}. Results may have expired or failed.', 'error')
             return redirect(url_for('index'))
        else:
             print(f"   📂 Dir Contents: {os.listdir(job_output_path)}")
             flash(f'Download file ({filename}) not found. It might not have been generated due to errors or no relevant data.', 'warning')
             return redirect(url_for('show_results', job_id=job_id))
    try:
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        print(f"   ❌ Error sending file: {e}");
        flash(f"Download error: {e}", 'error');
        return redirect(url_for('show_results', job_id=job_id))

# --- Run the App ---
if __name__ == '__main__':
    print("🚀 Starting Flask development server...")
    if not get_gemini_api_key_state():
        print("\n⚠️⚠️⚠️ WARNING: GEMINI_API_KEY not found or invalid. Gemini functionality will be disabled. ⚠️⚠️⚠️\n")
    else:
         print("✅ Gemini API Key detected.")
    app.run(host='127.0.0.1', port=5000, debug=True)

# --- End of app.py ---